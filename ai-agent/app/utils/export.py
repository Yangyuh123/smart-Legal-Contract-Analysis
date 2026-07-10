"""导出工具 - 将Markdown或文本转换为DOCX文档."""
from __future__ import annotations

import logging
import re
from typing import List, Optional, Dict, Any
from pathlib import Path
from io import BytesIO

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

logger = logging.getLogger(__name__)


class DocxExporter:
    """将法律文档内容导出为格式化的 .docx 文件."""

    def __init__(
        self,
        font_name: str = "仿宋",
        font_size: float = 12.0,
        title_font_name: str = "黑体",
        title_font_size: float = 16.0,
    ):
        """
        Args:
            font_name: 正文字体
            font_size: 正文字号（pt）
            title_font_name: 标题字体
            title_font_size: 标题字号（pt）
        """
        self.font_name = font_name
        self.font_size = font_size
        self.title_font_name = title_font_name
        self.title_font_size = title_font_size

    def export_markdown(
        self,
        markdown_text: str,
        output_path: Optional[str] = None,
        title: Optional[str] = None,
    ) -> bytes:
        """
        将Markdown文本导出为DOCX。

        Args:
            markdown_text: Markdown格式文本
            output_path: 输出文件路径（可选，不提供则返回bytes）
            title: 文档标题（可选，用于首页标题）

        Returns:
            DOCX文件的bytes内容
        """
        doc = Document()

        # 设置默认样式
        self._setup_styles(doc)

        # 添加文档标题
        if title:
            title_para = doc.add_paragraph()
            title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            title_run = title_para.add_run(title)
            title_run.font.name = self.title_font_name
            title_run.font.size = Pt(self.title_font_size)
            title_run.bold = True
            self._set_run_font(title_run, self.title_font_name)

            # 标题下空行
            doc.add_paragraph()

        # 逐行解析Markdown
        lines = markdown_text.split("\n")
        i = 0
        while i < len(lines):
            line = lines[i]

            # 空行
            if not line.strip():
                i += 1
                continue

            # 标题行
            heading_match = re.match(r"^(#{1,6})\s+(.+)$", line)
            if heading_match:
                level = len(heading_match.group(1))
                text = heading_match.group(2).strip()
                heading = doc.add_heading(text, level=min(level, 3))
                self._set_run_font_in_para(heading, self.title_font_name)
                i += 1
                continue

            # 无序列表
            list_match = re.match(r"^[\-\*\+]\s+(.+)$", line)
            if list_match:
                para = doc.add_paragraph(list_match.group(1).strip(), style="List Bullet")
                self._set_run_font_in_para(para, self.font_name)
                i += 1
                continue

            # 有序列表
            ordered_match = re.match(r"^\d+[\.\)、]\s+(.+)$", line)
            if ordered_match:
                para = doc.add_paragraph(ordered_match.group(1).strip(), style="List Number")
                self._set_run_font_in_para(para, self.font_name)
                i += 1
                continue

            # 引用块
            blockquote_match = re.match(r"^>\s*(.+)$", line)
            if blockquote_match:
                para = doc.add_paragraph()
                para.paragraph_format.left_indent = Cm(1.5)
                run = para.add_run(blockquote_match.group(1).strip())
                run.font.name = self.font_name
                run.font.size = Pt(self.font_size)
                run.italic = True
                run.font.color.rgb = RGBColor(100, 100, 100)
                self._set_run_font(run, self.font_name)
                i += 1
                continue

            # 代码块
            if line.strip().startswith("```"):
                code_lines = []
                i += 1
                while i < len(lines) and not lines[i].strip().startswith("```"):
                    code_lines.append(lines[i])
                    i += 1
                if code_lines:
                    code_text = "\n".join(code_lines)
                    para = doc.add_paragraph()
                    run = para.add_run(code_text)
                    run.font.name = "Consolas"
                    run.font.size = Pt(9)
                    run.font.color.rgb = RGBColor(50, 50, 50)
                    self._set_run_font(run, "Consolas")
                i += 1
                continue

            # 分隔线
            if re.match(r"^[\-\*_]{3,}$", line.strip()):
                doc.add_paragraph("_" * 40)
                i += 1
                continue

            # 普通段落
            para = doc.add_paragraph()
            # 处理行内标记
            self._parse_inline_markdown(para, line, self.font_name, self.font_size)
            i += 1

        # 输出
        if output_path:
            output_path = str(Path(output_path).resolve())
            doc.save(output_path)
            logger.info(f"DOCX exported to: {output_path}")
            with open(output_path, "rb") as f:
                return f.read()

        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.read()

    def export_plain_text(
        self,
        text: str,
        output_path: Optional[str] = None,
        title: Optional[str] = None,
    ) -> bytes:
        """
        将纯文本导出为DOCX。

        Args:
            text: 纯文本
            output_path: 输出文件路径
            title: 文档标题

        Returns:
            DOCX bytes
        """
        doc = Document()
        self._setup_styles(doc)

        # 标题
        if title:
            title_para = doc.add_paragraph()
            title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            title_run = title_para.add_run(title)
            title_run.font.name = self.title_font_name
            title_run.font.size = Pt(self.title_font_size)
            title_run.bold = True
            self._set_run_font(title_run, self.title_font_name)
            doc.add_paragraph()

        # 按段落添加
        paragraphs = text.split("\n")
        for para_text in paragraphs:
            para_text = para_text.strip()
            if para_text:
                para = doc.add_paragraph()
                run = para.add_run(para_text)
                run.font.name = self.font_name
                run.font.size = Pt(self.font_size)
                self._set_run_font(run, self.font_name)

        if output_path:
            output_path = str(Path(output_path).resolve())
            doc.save(output_path)
            logger.info(f"DOCX exported to: {output_path}")
            with open(output_path, "rb") as f:
                return f.read()

        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.read()

    def export_from_template(
        self,
        template_path: str,
        context: Dict[str, Any],
        output_path: Optional[str] = None,
    ) -> bytes:
        """
        使用Jinja2模板生成DOCX。

        Args:
            template_path: DOCX模板文件路径
            context: 模板变量上下文
            output_path: 输出路径

        Returns:
            DOCX bytes

        Note:
            对于复杂的模板填充，建议使用 docxtpl 库。
            这里提供基础实现。
        """
        from jinja2 import Template

        doc = Document(template_path)

        for para in doc.paragraphs:
            if "{{" in para.text and "}}" in para.text:
                try:
                    template = Template(para.text)
                    rendered = template.render(**context)
                    # 清空并重新设置
                    for run in para.runs:
                        run.text = ""
                    para.runs[0].text = rendered if para.runs else rendered
                except Exception as e:
                    logger.warning(f"Template rendering failed for paragraph: {e}")

        if output_path:
            doc.save(output_path)
            with open(output_path, "rb") as f:
                return f.read()

        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.read()

    def _setup_styles(self, doc: Document):
        """设置文档默认样式."""
        style = doc.styles["Normal"]
        style.font.name = self.font_name
        style.font.size = Pt(self.font_size)
        style.paragraph_format.line_spacing = 1.5
        style.paragraph_format.space_after = Pt(6)
        self._set_run_font_in_para(style, self.font_name)

        # 设置标题样式
        for level in range(1, 4):
            heading_style = doc.styles[f"Heading {level}"]
            heading_style.font.name = self.title_font_name
            heading_style.font.color.rgb = RGBColor(0, 0, 0)
            self._set_run_font_in_para(heading_style, self.title_font_name)

    def _parse_inline_markdown(self, para, text: str, font_name: str, font_size: float):
        """
        解析行内Markdown格式（粗体、斜体等）。

        Args:
            para: docx Paragraph对象
            text: 原始文本行
            font_name: 默认字体
            font_size: 默认字号
        """
        # 简化处理：识别 **粗体** 标记
        pattern = re.compile(r"(\*\*(.+?)\*\*)|(\*(.+?)\*)|(___(.+?)___)|(\[(.+?)\]\((.+?)\))")
        last_end = 0

        for match in pattern.finditer(text):
            # 前面的普通文本
            normal_text = text[last_end:match.start()]
            if normal_text:
                run = para.add_run(normal_text)
                run.font.name = font_name
                run.font.size = Pt(font_size)
                self._set_run_font(run, font_name)

            # 粗体
            if match.group(1):
                run = para.add_run(match.group(2))
                run.bold = True
                run.font.name = font_name
                run.font.size = Pt(font_size)
                self._set_run_font(run, font_name)
            # 斜体
            elif match.group(3):
                run = para.add_run(match.group(4))
                run.italic = True
                run.font.name = font_name
                run.font.size = Pt(font_size)
                self._set_run_font(run, font_name)
            # 链接 (显示为文本)
            elif match.group(7):
                run = para.add_run(f"{match.group(8)} ({match.group(9)})")
                run.font.color.rgb = RGBColor(0, 0, 255)
                run.underline = True
                run.font.name = font_name
                run.font.size = Pt(font_size)
                self._set_run_font(run, font_name)

            last_end = match.end()

        # 剩余文本
        remaining = text[last_end:]
        if remaining:
            run = para.add_run(remaining)
            run.font.name = font_name
            run.font.size = Pt(font_size)
            self._set_run_font(run, font_name)

    def _set_run_font(self, run, font_name: str):
        """设置run的字体，处理中文字体."""
        run.font.name = font_name
        run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)

    def _set_run_font_in_para(self, para, font_name: str):
        """设置段落中所有run的字体."""
        for run in para.runs:
            self._set_run_font(run, font_name)
        # 也设置段落级别的字体
        if hasattr(para, "style") and para.style:
            try:
                para.style.font.name = font_name
            except Exception:
                pass
