"""DOCX文档解析器 - 使用python-docx提取Word文档的结构化内容."""
from __future__ import annotations

import logging
import re
from typing import List, Dict, Any, Optional
from pathlib import Path

from docx import Document

from app.parsers.text_cleaner import clean_contract_text, normalize_clause_text

logger = logging.getLogger(__name__)


class DocxParser:
    """解析 .docx 格式的Word文档，提取文本、表格和结构信息."""

    def __init__(self, extract_styles: bool = True, extract_tables: bool = True):
        """
        Args:
            extract_styles: 是否提取样式信息（用于识别标题/正文）
            extract_tables: 是否提取表格内容
        """
        self.extract_styles = extract_styles
        self.extract_tables = extract_tables

    def parse(self, file_path: str) -> Dict[str, Any]:
        """
        解析DOCX文件。

        Args:
            file_path: DOCX文件路径

        Returns:
            解析结果字典，包含:
            - full_text: 完整文本
            - paragraphs: 段落列表
            - tables: 表格列表
            - structure: 文档结构（标题层级）
            - metadata: 文档元数据

        Raises:
            FileNotFoundError: 文件不存在
            ValueError: 文件格式错误
        """
        path = Path(file_path)
        if not path.exists():
            raise FileNotFoundError(f"文件不存在: {file_path}")

        if path.suffix.lower() not in (".docx", ".doc"):
            raise ValueError(f"不支持的文件格式: {path.suffix}，仅支持 .docx")

        try:
            doc = Document(str(path))
        except Exception as e:
            raise ValueError(f"无法打开文档 {file_path}: {e}")

        result: Dict[str, Any] = {
            "file_name": path.name,
            "paragraphs": [],
            "tables": [],
            "structure": [],
            "metadata": self._extract_metadata(doc),
            "full_text": "",
        }

        # 提取段落
        paragraphs_data = []
        for para in doc.paragraphs:
            para_data = self._parse_paragraph(para)
            paragraphs_data.append(para_data)
            result["paragraphs"].append(para_data)

        # 提取表格
        if self.extract_tables:
            tables_data = []
            for idx, table in enumerate(doc.tables):
                table_data = self._parse_table(table, idx)
                tables_data.append(table_data)
                result["tables"].append(table_data)

        # 构建结构化文本
        text_parts = []
        current_section = None

        for para_data in paragraphs_data:
            text = para_data["text"]
            if not text:
                continue

            style = para_data.get("style_name", "")

            # 识别标题
            if para_data.get("is_heading"):
                level = para_data.get("heading_level", 1)
                result["structure"].append({
                    "type": "heading",
                    "level": level,
                    "text": text,
                })
                text_parts.append(("heading", level, text))
                current_section = text
            else:
                text_parts.append(("body", 0, text))

        # 构建全文
        full_lines = []
        for part in text_parts:
            typ, level, text = part
            if typ == "heading":
                full_lines.append("")  # 标题前空行
                prefix = "#" * min(level, 6)
                full_lines.append(f"{prefix} {text}")
                full_lines.append("")
            else:
                full_lines.append(text)

        raw_text = "\n".join(full_lines)

        # 清洗文本
        result["full_text"] = clean_contract_text(raw_text)

        return result

    def _parse_paragraph(self, para) -> Dict[str, Any]:
        """解析单个段落."""
        text = para.text.strip() if para.text else ""

        para_data: Dict[str, Any] = {
            "text": text,
            "style_name": para.style.name if para.style else "",
            "is_heading": False,
            "heading_level": 0,
        }

        # 检查是否为标题
        style_name = para.style.name.lower() if para.style else ""
        if "heading" in style_name or "标题" in style_name:
            para_data["is_heading"] = True
            # 尝试提取标题级别
            level_match = re.search(r"heading\s*(\d+)", style_name, re.IGNORECASE)
            if level_match:
                para_data["heading_level"] = int(level_match.group(1))
            elif "标题" in style_name:
                level_match = re.search(r"标题\s*(\d+)", style_name)
                if level_match:
                    para_data["heading_level"] = int(level_match.group(1))
                else:
                    para_data["heading_level"] = 1
            else:
                para_data["heading_level"] = 1

        # 提取样式信息
        if self.extract_styles and para.style:
            para_data["bold"] = para.style.font.bold if para.style.font else None
            para_data["italic"] = para.style.font.italic if para.style.font else None
            para_data["font_size"] = str(para.style.font.size) if para.style.font and para.style.font.size else None

        # 提取 run 级别的格式信息
        runs_data = []
        for run in para.runs:
            run_info = {"text": run.text}
            if run.bold:
                run_info["bold"] = True
            if run.italic:
                run_info["italic"] = True
            if run.underline:
                run_info["underline"] = True
            runs_data.append(run_info)
        para_data["runs"] = runs_data

        return para_data

    def _parse_table(self, table, table_idx: int) -> Dict[str, Any]:
        """解析表格."""
        headers = []
        rows = []

        for row_idx, row in enumerate(table.rows):
            cells = [cell.text.strip() for cell in row.cells]
            if row_idx == 0:
                headers = cells
            else:
                rows.append(cells)

        return {
            "table_index": table_idx,
            "headers": headers,
            "rows": rows,
            "row_count": len(rows),
            "column_count": len(headers),
        }

    def _extract_metadata(self, doc: Document) -> Dict[str, Any]:
        """提取文档元数据."""
        props = doc.core_properties
        return {
            "author": props.author or "",
            "created": str(props.created) if props.created else "",
            "modified": str(props.modified) if props.modified else "",
            "last_modified_by": props.last_modified_by or "",
            "title": props.title or "",
            "subject": props.subject or "",
            "revision": props.revision or "",
        }

    def parse_text_only(self, file_path: str) -> str:
        """
        仅提取文本内容（快速模式）。

        Args:
            file_path: DOCX文件路径

        Returns:
            清洗后的纯文本
        """
        path = Path(file_path)
        if not path.exists():
            raise FileNotFoundError(f"文件不存在: {file_path}")

        doc = Document(str(path))
        paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]

        # 也提取表格中的文本
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        paragraphs.append(cell.text.strip())

        text = "\n".join(paragraphs)
        return clean_contract_text(text)

    def extract_clauses(self, file_path: str) -> List[Dict[str, Any]]:
        """
        从合同文档中提取条款列表。

        Args:
            file_path: DOCX 文件路径

        Returns:
            条款列表 [{"clause_no": "第一条", "title": "...", "content": "..."}, ...]
        """
        from app.parsers.text_cleaner import extract_clauses

        text = self.parse_text_only(file_path)
        return extract_clauses(text)
